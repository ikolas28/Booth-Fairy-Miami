from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import KeepTogether, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


OUTPUT_DIR = Path("crm/contracts/outputs")
PUBLIC_DIR = Path("website/assets/contracts")

EN_DOCX = OUTPUT_DIR / "Booth Fairy Miami Service Agreement - English.docx"
ES_DOCX = OUTPUT_DIR / "Booth Fairy Miami Acuerdo de Servicios - Espanol.docx"
EN_PDF = PUBLIC_DIR / "booth-fairy-miami-service-agreement-english.pdf"
ES_PDF = PUBLIC_DIR / "booth-fairy-miami-acuerdo-de-servicios-espanol.pdf"


ENGLISH = {
    "title": "Booth Fairy Miami Service Agreement",
    "subtitle": "DSLR print photo booth, instant prints, instant digital sharing, and premium DJ services",
    "client_info": "Client Information",
    "event_details": "Event Details",
    "services_selected": "Services Selected",
    "signatures": "Signatures",
    "field_rows": {
        "client": [
            ("Client name", ""),
            ("Company name, if applicable", ""),
            ("Billing address", ""),
            ("Phone", ""),
            ("Email", ""),
        ],
        "event": [
            ("Event date", ""),
            ("Event type", ""),
            ("Start time", ""),
            ("End time", ""),
            ("Venue name", ""),
            ("Venue address", ""),
            ("City", ""),
            ("Venue contact", ""),
            ("Venue phone/email", ""),
            ("Guest count", ""),
            ("Special instructions", ""),
        ],
        "services": [
            ("DSLR print photo booth", "Included / Not included"),
            ("Premium DJ services", "Included / Not included"),
            ("Package or custom quote", ""),
            ("Total service fee", "$"),
            ("Retainer/deposit due", "50% due at contract signing to reserve the event date"),
            ("Balance due date", "Remaining 50% due on the day of the event"),
            ("Payment link or invoice", ""),
        ],
        "signatures": [
            ("Client signature", ""),
            ("Client printed name", ""),
            ("Date", ""),
            ("Company representative", ""),
            ("Company signature", ""),
            ("Date", ""),
        ],
    },
    "clauses": [
        ("DSLR Print Photo Booth", [
            "Every DSLR Print Photo Booth package includes instant prints, instant digital sharing, a premium backdrop, studio flash lighting, custom event overlay, props, an attendant, unlimited sessions, and digital gallery delivery.",
            "Pricing: 2 hours is $450, 3 hours is $575, and 4 hours is $700. Additional time, travel, parking, venue fees, premium add-ons, or DJ services must be confirmed in writing and may increase the total service fee.",
        ]),
        ("Agreement", [
            "This Service Agreement is between the client listed above (Client) and Booth Fairy Miami (Company). Company agrees to provide the event services selected in this Agreement and any attached written quote or invoice. This Agreement, together with any written quote, invoice, or approved addendum, is the entire agreement between the parties. Any change must be confirmed in writing by both parties.",
            "Company provides DSLR print photo booth services with instant prints, digital gallery delivery, and instant digital sharing. Company does not provide 360 photo booth services under this Agreement. Company also provides premium DJ services when selected in the Services Selected section or in an attached written quote.",
        ]),
        ("Service Period", [
            "Company will provide the selected services during the service period listed in the Event Details. Setup and breakdown time are not included in the paid service period unless expressly stated in the written quote.",
            "Additional event time must be requested by Client and approved by Company. Overtime is subject to Company availability and will be billed at the overtime rate stated in the quote or invoice, or at a rate mutually confirmed in writing before overtime begins.",
        ]),
        ("DSLR Print Photo Booth Services", [
            "When DSLR print photo booth services are selected, Company will provide a professional DSLR photo booth setup designed to capture high-quality photographs. The experience includes instant prints during the booked service period, studio flash lighting, camera, booth station, attendant, digital gallery, and instant digital sharing.",
            "Digital sharing may include text, email, QR code, online gallery, or similar delivery methods, subject to venue internet access, mobile carrier availability, device compatibility, and third-party platform performance. Company will make commercially reasonable efforts to support instant digital sharing during the event, but cannot guarantee uninterrupted wireless or internet-dependent delivery.",
            "Client understands that instant prints are included during the booked service period. A 360 photo booth service is not included.",
        ]),
        ("Premium DJ Services", [
            "When DJ services are selected, Company will provide professional DJ entertainment for the service period listed in the Agreement or quote. DJ services may include music playback, basic event announcements, sound equipment, microphones, and lighting elements depending on the package selected.",
            "Client is responsible for providing any required timeline, music preferences, must-play songs, do-not-play songs, pronunciation notes, special dances, announcements, and event formalities in advance.",
        ]),
        ("Retainer, Payment, and Booking Confirmation", [
            "A booking is not confirmed until Company has checked availability, Client has signed this Agreement or approved the written quote, and the required retainer/deposit or payment has been received and confirmed by Company.",
            "A non-refundable retainer/deposit equal to 50% of the total service fee is due when Client signs the contract. The retainer/deposit reserves the event date, time, staff, and equipment and is non-refundable to the extent permitted by applicable law because Company may decline other booking opportunities after reserving the date.",
            "The remaining 50% balance is due on the day of the event unless the written quote or invoice states an earlier due date.",
        ]),
        ("Changes, Rescheduling, and Cancellations", [
            "Any request to change the date, time, venue, service package, or scope of services must be made in writing. Changes are subject to Company availability, staffing, equipment availability, venue requirements, and any additional fees that apply.",
            "If Client requests to reschedule and Company is available for the new date and time, Company may apply the retainer/deposit to the new booking at Company's discretion and subject to a written updated agreement or invoice. If Company is not available, the original cancellation terms apply.",
        ]),
        ("Venue Access, Parking, Space, and Power", [
            "Client is responsible for ensuring timely venue access for setup and breakdown. Client must provide accurate venue details, load-in instructions, floor level, elevator access, parking information, and vendor rules before the event.",
            "Client is responsible for parking fees, valet fees, loading dock fees, vendor permits, venue-required insurance charges, and access fees charged by the venue or property.",
        ]),
        ("Outdoor Events and Weather", [
            "Outdoor service is subject to safe weather conditions, adequate covering, level ground, power safety, and venue approval. Client must provide weather protection from rain, wind, direct water exposure, excessive heat, and unsafe conditions.",
            "Company may pause, relocate, delay, or end service if weather, electrical conditions, crowd behavior, venue conditions, or safety concerns could damage equipment or create risk.",
        ]),
        ("Digital Gallery and Image Delivery", [
            "For DSLR print photo booth bookings, Company will also provide digital photo access through an online gallery, downloadable link, sharing platform, or comparable digital delivery method. Delivery timing may vary based on event size, internet access, file processing, platform availability, and package selected.",
            "Client is responsible for downloading, saving, backing up, and archiving delivered files. Company may retain files for a limited period but does not guarantee permanent storage or indefinite gallery availability.",
        ]),
        ("Client Content and Approvals", [
            "Client is responsible for providing logos, brand assets, event hashtags, names, dates, monograms, inspiration, music requests, timeline details, and required approvals by the deadline requested by Company.",
            "If a digital overlay, gallery branding, or event-specific design is included, Company may provide a draft and reasonable revisions as stated in the quote. Additional revision rounds, rush changes, or changes requested after approval may require additional fees.",
        ]),
        ("Model Release and Marketing Use", [
            "Unless Client opts out in writing before the event, Client grants Company permission to use event photos, videos, behind-the-scenes content, setup photos, and guest interaction content captured in connection with the services for Company portfolio, website, social media, advertising, and other good-faith business promotion.",
        ]),
        ("Guest Conduct and Equipment Protection", [
            "Client is responsible for the conduct of Client's guests, invitees, vendors, and venue representatives. Company may pause or stop service if guests misuse equipment, interfere with staff, create unsafe conditions, or violate venue rules.",
            "Client is responsible for damage, loss, or excessive cleaning caused by Client, guests, venue staff, or third parties, except to the extent caused by Company's gross negligence or willful misconduct.",
        ]),
        ("Limitation of Liability", [
            "Company will make commercially reasonable efforts to provide the selected services in a professional manner. If Company is unable to provide a substantial portion of the selected service because of Company equipment failure or Company fault, Company's maximum liability will be limited to a prorated refund of the affected service fee.",
            "Company is not responsible for service interruptions or losses caused by venue restrictions, power failure, internet failure, mobile network issues, guest conduct, weather, acts of God, government orders, emergencies, road closures, accidents, illness, or other circumstances outside Company's reasonable control.",
        ]),
        ("Governing Law", [
            "This Agreement is governed by the laws of the State of Florida. Venue for any dispute will be in the county of Company's principal place of business unless applicable law requires otherwise.",
        ]),
        ("Final Agreement and Signatures", [
            "Client has read and understands this Agreement. By signing below, Client confirms that all event details are accurate to the best of Client's knowledge and agrees to the terms above.",
        ]),
    ],
}


SPANISH = {
    "title": "Acuerdo de Servicios de Booth Fairy Miami",
    "subtitle": "Cabina fotografica DSLR con impresiones instantaneas, envio digital instantaneo y servicios premium de DJ",
    "client_info": "Informacion del Cliente",
    "event_details": "Detalles del Evento",
    "services_selected": "Servicios Seleccionados",
    "signatures": "Firmas",
    "field_rows": {
        "client": [
            ("Nombre del cliente", ""),
            ("Nombre de la compania, si aplica", ""),
            ("Direccion de facturacion", ""),
            ("Telefono", ""),
            ("Correo electronico", ""),
        ],
        "event": [
            ("Fecha del evento", ""),
            ("Tipo de evento", ""),
            ("Hora de inicio", ""),
            ("Hora de finalizacion", ""),
            ("Nombre del lugar", ""),
            ("Direccion del lugar", ""),
            ("Ciudad", ""),
            ("Contacto del lugar", ""),
            ("Telefono/correo del lugar", ""),
            ("Cantidad de invitados", ""),
            ("Instrucciones especiales", ""),
        ],
        "services": [
            ("Cabina fotografica DSLR con impresiones", "Incluido / No incluido"),
            ("Servicios premium de DJ", "Incluido / No incluido"),
            ("Paquete o cotizacion personalizada", ""),
            ("Tarifa total del servicio", "$"),
            ("Retenedor/deposito requerido", "50% al firmar el contrato para reservar la fecha"),
            ("Fecha de pago del balance", "50% restante el dia del evento"),
            ("Enlace de pago o factura", ""),
        ],
        "signatures": [
            ("Firma del cliente", ""),
            ("Nombre impreso del cliente", ""),
            ("Fecha", ""),
            ("Representante de la compania", ""),
            ("Firma de la compania", ""),
            ("Fecha", ""),
        ],
    },
    "clauses": [
        ("Cabina Fotografica DSLR con Impresiones", [
            "Cada paquete incluye impresiones instantaneas, envio digital instantaneo, fondo premium, iluminacion de flash de estudio, overlay personalizado, accesorios, asistente, sesiones ilimitadas y entrega de galeria digital.",
            "Precios: 2 horas $450, 3 horas $575 y 4 horas $700. Tiempo adicional, viajes, estacionamiento, cargos del lugar, servicios adicionales premium o servicios de DJ deben confirmarse por escrito y pueden aumentar la tarifa total.",
        ]),
        ("Acuerdo", [
            "Este Acuerdo de Servicios es entre el cliente indicado arriba (Cliente) y Booth Fairy Miami (Compania). La Compania acepta proporcionar los servicios seleccionados en este Acuerdo y en cualquier cotizacion o factura escrita adjunta. Este Acuerdo, junto con cualquier cotizacion, factura o anexo aprobado por escrito, constituye el acuerdo completo entre las partes. Cualquier cambio debe confirmarse por escrito por ambas partes.",
            "La Compania ofrece servicios de cabina fotografica DSLR con impresiones instantaneas, entrega de galeria digital y envio digital instantaneo. La Compania no ofrece servicio de cabina 360 bajo este Acuerdo. La Compania tambien ofrece servicios premium de DJ cuando sean seleccionados en este Acuerdo o en una cotizacion escrita.",
        ]),
        ("Periodo de Servicio", [
            "La Compania proporcionara los servicios seleccionados durante el periodo indicado en los Detalles del Evento. El tiempo de montaje y desmontaje no esta incluido en el periodo pagado del servicio, a menos que se indique expresamente en la cotizacion escrita.",
            "Cualquier tiempo adicional debe ser solicitado por el Cliente y aprobado por la Compania. El tiempo extra esta sujeto a disponibilidad y se cobrara segun la tarifa indicada en la cotizacion o factura, o segun una tarifa confirmada por escrito antes de comenzar.",
        ]),
        ("Servicios de Cabina Fotografica Digital DSLR", [
            "Cuando se seleccionan servicios de cabina fotografica DSLR con impresiones, la Compania proporcionara una configuracion profesional disenada para capturar fotografias de alta calidad. La experiencia incluye impresiones instantaneas durante el periodo contratado, iluminacion, camara, estacion de cabina, asistente, galeria digital y envio digital instantaneo.",
            "El envio digital puede incluir texto, correo electronico, codigo QR, galeria en linea u otros metodos similares, sujeto al acceso de internet del lugar, senal celular, compatibilidad de dispositivos y plataformas de terceros. La Compania hara esfuerzos comercialmente razonables para apoyar el envio digital instantaneo durante el evento, pero no garantiza servicio ininterrumpido dependiente de internet o senal inalambrica.",
            "El Cliente entiende que las impresiones instantaneas estan incluidas durante el periodo contratado. El servicio de cabina 360 no esta incluido.",
        ]),
        ("Servicios Premium de DJ", [
            "Cuando se seleccionan servicios de DJ, la Compania proporcionara entretenimiento profesional de DJ durante el periodo indicado en el Acuerdo o cotizacion. Los servicios pueden incluir reproduccion de musica, anuncios basicos, equipo de sonido, microfonos y elementos de iluminacion segun el paquete seleccionado.",
            "El Cliente es responsable de proporcionar con anticipacion el itinerario, preferencias musicales, canciones obligatorias, canciones prohibidas, notas de pronunciacion, bailes especiales, anuncios y formalidades del evento.",
        ]),
        ("Retenedor, Pago y Confirmacion de Reserva", [
            "Una reserva no esta confirmada hasta que la Compania haya verificado disponibilidad, el Cliente haya firmado este Acuerdo o aprobado la cotizacion escrita, y la Compania haya recibido y confirmado el retenedor/deposito o pago requerido.",
            "Un retenedor/deposito no reembolsable equivalente al 50% de la tarifa total del servicio vence al firmar el contrato. El retenedor/deposito reserva la fecha, horario, personal y equipo, y no es reembolsable en la medida permitida por la ley aplicable porque la Compania puede rechazar otras oportunidades de reserva al apartar la fecha.",
            "El 50% restante vence el dia del evento, a menos que la cotizacion o factura escrita indique una fecha anterior.",
        ]),
        ("Cambios, Reprogramaciones y Cancelaciones", [
            "Cualquier solicitud para cambiar fecha, hora, lugar, paquete o alcance de servicios debe hacerse por escrito. Los cambios estan sujetos a disponibilidad, personal, equipo, requisitos del lugar y cargos adicionales aplicables.",
            "Si el Cliente solicita reprogramar y la Compania esta disponible para la nueva fecha y hora, la Compania puede aplicar el retenedor/deposito a la nueva reserva a discrecion de la Compania y sujeto a un acuerdo o factura actualizada por escrito. Si la Compania no esta disponible, aplican los terminos originales de cancelacion.",
        ]),
        ("Acceso al Lugar, Estacionamiento, Espacio y Energia", [
            "El Cliente es responsable de asegurar acceso oportuno al lugar para montaje y desmontaje. El Cliente debe proporcionar detalles precisos del lugar, instrucciones de entrada, nivel de piso, acceso a elevador, informacion de estacionamiento y reglas para proveedores antes del evento.",
            "El Cliente es responsable por cargos de estacionamiento, valet, muelle de carga, permisos de proveedor, cargos de seguro requeridos por el lugar y cualquier cargo de acceso cobrado por el lugar o propiedad.",
        ]),
        ("Eventos al Aire Libre y Clima", [
            "El servicio al aire libre esta sujeto a condiciones climaticas seguras, cobertura adecuada, terreno nivelado, seguridad electrica y aprobacion del lugar. El Cliente debe proporcionar proteccion contra lluvia, viento, agua directa, calor excesivo y condiciones inseguras.",
            "La Compania puede pausar, reubicar, retrasar o terminar el servicio si el clima, condiciones electricas, comportamiento de invitados, condiciones del lugar o temas de seguridad pueden danar el equipo o crear riesgo.",
        ]),
        ("Galeria Digital y Entrega de Imagenes", [
            "Para reservas de cabina fotografica DSLR con impresiones, la Compania tambien proporcionara acceso digital a las fotos mediante galeria en linea, enlace de descarga, plataforma de envio u otro metodo comparable. El tiempo de entrega puede variar segun el tamano del evento, internet, procesamiento de archivos, disponibilidad de plataformas y paquete seleccionado.",
            "El Cliente es responsable de descargar, guardar, respaldar y archivar los archivos entregados. La Compania puede conservar archivos por un periodo limitado, pero no garantiza almacenamiento permanente ni disponibilidad indefinida de la galeria.",
        ]),
        ("Contenido del Cliente y Aprobaciones", [
            "El Cliente es responsable de proporcionar logos, recursos de marca, hashtags, nombres, fechas, monogramas, inspiracion, solicitudes de musica, itinerario y aprobaciones requeridas antes de la fecha limite solicitada por la Compania.",
            "Si se incluye overlay digital, marca de galeria o diseno especifico del evento, la Compania puede proporcionar un borrador y revisiones razonables segun la cotizacion. Rondas adicionales, cambios urgentes o cambios solicitados despues de aprobar pueden requerir cargos adicionales.",
        ]),
        ("Autorizacion de Imagen y Uso de Marketing", [
            "A menos que el Cliente solicite por escrito no participar antes del evento, el Cliente otorga permiso a la Compania para usar fotos, videos, contenido detras de camaras, fotos del montaje e interacciones de invitados capturadas en relacion con los servicios para portafolio, sitio web, redes sociales, publicidad y promocion comercial de buena fe.",
        ]),
        ("Conducta de Invitados y Proteccion del Equipo", [
            "El Cliente es responsable por la conducta de sus invitados, representantes, proveedores y personal del lugar. La Compania puede pausar o detener el servicio si los invitados mal utilizan el equipo, interfieren con el personal, crean condiciones inseguras o violan reglas del lugar.",
            "El Cliente es responsable por dano, perdida o limpieza excesiva causada por el Cliente, invitados, personal del lugar o terceros, excepto en la medida causada por negligencia grave o conducta intencional de la Compania.",
        ]),
        ("Limitacion de Responsabilidad", [
            "La Compania hara esfuerzos comercialmente razonables para proporcionar los servicios seleccionados de manera profesional. Si la Compania no puede proporcionar una parte sustancial del servicio seleccionado debido a falla de equipo de la Compania o culpa de la Compania, la responsabilidad maxima de la Compania se limita a un reembolso prorrateado de la tarifa del servicio afectado.",
            "La Compania no es responsable por interrupciones o perdidas causadas por restricciones del lugar, fallas electricas, fallas de internet, problemas de red movil, conducta de invitados, clima, actos de Dios, ordenes gubernamentales, emergencias, cierres de carreteras, accidentes, enfermedad u otras circunstancias fuera del control razonable de la Compania.",
        ]),
        ("Ley Aplicable", [
            "Este Acuerdo se rige por las leyes del Estado de Florida. El lugar para cualquier disputa sera el condado del lugar principal de negocios de la Compania, salvo que la ley aplicable requiera otra cosa.",
        ]),
        ("Acuerdo Final y Firmas", [
            "El Cliente ha leido y entiende este Acuerdo. Al firmar abajo, el Cliente confirma que todos los detalles del evento son exactos segun su mejor conocimiento y acepta los terminos anteriores.",
        ]),
    ],
}


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(9.6)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, before, after in [("Heading 1", 14, 14, 4), ("Heading 2", 11, 8, 3)]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(37, 23, 45)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc, data):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(data["title"])
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(37, 23, 45)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(9)
    run = p.add_run(data["subtitle"])
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(9.4)
    run.font.color.rgb = RGBColor(85, 85, 85)


def add_field_table(doc, title, rows):
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        for cell in table.rows[i].cells:
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
                    run.font.size = Pt(8.8)
        for run in table.rows[i].cells[0].paragraphs[0].runs:
            run.font.bold = True


def add_clause(doc, heading, paragraphs):
    doc.add_heading(heading, level=1)
    for text in paragraphs:
        doc.add_paragraph(text)


def build_docx(data, path):
    doc = Document()
    style_doc(doc)
    add_title(doc, data)
    add_field_table(doc, data["client_info"], data["field_rows"]["client"])
    add_field_table(doc, data["event_details"], data["field_rows"]["event"])
    add_field_table(doc, data["services_selected"], data["field_rows"]["services"])
    for heading, paragraphs in data["clauses"]:
        add_clause(doc, heading, paragraphs)
    add_field_table(doc, data["signatures"], data["field_rows"]["signatures"])
    doc.save(path)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        "ContractTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=23,
        textColor=colors.HexColor("#25172d"),
        spaceAfter=5,
        alignment=0,
    ))
    styles.add(ParagraphStyle(
        "ContractSubtitle",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#555555"),
        spaceAfter=12,
    ))
    styles.add(ParagraphStyle(
        "ContractHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11.5,
        leading=14,
        textColor=colors.HexColor("#25172d"),
        spaceBefore=8,
        spaceAfter=3,
    ))
    styles.add(ParagraphStyle(
        "ContractBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8.4,
        leading=10.2,
        spaceAfter=3,
    ))
    return styles


def add_pdf_table(story, title, rows, styles):
    story.append(Paragraph(title, styles["ContractHeading"]))
    table = Table(rows, colWidths=[2.25 * inch, 4.55 * inch], hAlign="LEFT")
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#d8cbdc")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#fff4f7")),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("LEADING", (0, 0), (-1, -1), 10),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(table)
    story.append(Spacer(1, 7))


def add_pdf_signature_table(story, title, rows, styles):
    paired_rows = [
        [rows[0][0], rows[0][1], rows[4][0], rows[4][1]],
        [rows[1][0], rows[1][1], rows[3][0], rows[3][1]],
        [rows[2][0], rows[2][1], rows[5][0], rows[5][1]],
    ]
    heading = Paragraph(title, styles["ContractHeading"])
    table = Table(paired_rows, colWidths=[1.65 * inch, 1.65 * inch, 1.85 * inch, 1.65 * inch], hAlign="LEFT")
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#d8cbdc")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#fff4f7")),
        ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#fff4f7")),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("LEADING", (0, 0), (-1, -1), 10),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]))
    story.append(KeepTogether([heading, table, Spacer(1, 7)]))


def build_pdf(data, path):
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        rightMargin=0.65 * inch,
        leftMargin=0.65 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
        title=data["title"],
    )
    story = [
        Paragraph(data["title"], styles["ContractTitle"]),
        Paragraph(data["subtitle"], styles["ContractSubtitle"]),
    ]
    add_pdf_table(story, data["client_info"], data["field_rows"]["client"], styles)
    add_pdf_table(story, data["event_details"], data["field_rows"]["event"], styles)
    add_pdf_table(story, data["services_selected"], data["field_rows"]["services"], styles)
    for heading, paragraphs in data["clauses"]:
        story.append(Paragraph(heading, styles["ContractHeading"]))
        for text in paragraphs:
            story.append(Paragraph(text, styles["ContractBody"]))
    add_pdf_signature_table(story, data["signatures"], data["field_rows"]["signatures"], styles)
    doc.build(story)


def build():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    PUBLIC_DIR.mkdir(parents=True, exist_ok=True)
    build_docx(ENGLISH, EN_DOCX)
    build_docx(SPANISH, ES_DOCX)
    build_pdf(ENGLISH, EN_PDF)
    build_pdf(SPANISH, ES_PDF)


if __name__ == "__main__":
    build()
